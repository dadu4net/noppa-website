import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def hex_to_rgb(hex_code):
    hex_code = hex_code.lstrip('#')
    return tuple(int(hex_code[i:i+2], 16) for i in (0, 2, 4))

def setup_style(style, font_name, size_pt, color_hex, bold=False):
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    r, g, b = hex_to_rgb(color_hex)
    font.color.rgb = RGBColor(r, g, b)
    font.bold = bold

def create_template():
    doc = Document()
    styles = doc.styles

    # Colors
    INK = '#1A2440'
    NAVY = '#0F2A66'
    
    # 1. Normal (Body)
    setup_style(styles['Normal'], 'Kentledge', 10.5, INK, bold=False)
    
    # 2. Heading 1
    setup_style(styles['Heading 1'], 'Kentledge', 36, NAVY, bold=True)
    
    # 3. Heading 2
    setup_style(styles['Heading 2'], 'Kentledge', 22, NAVY, bold=True)
    
    # 4. Heading 3
    setup_style(styles['Heading 3'], 'Kentledge', 14, NAVY, bold=True)
    
    # Create sample content
    doc.add_heading('Noppa Basis Sjabloon', level=1)
    doc.add_paragraph('Dit sjabloon bevat de basisstijlen voor Noppa Solutions & Consultants, geconfigureerd volgens de Brand Guide v1.0.', style='Normal')
    
    doc.add_heading('Koppen & Typografie', level=2)
    doc.add_paragraph('Hieronder zie je voorbeelden van de verschillende hiërarchische niveaus.', style='Normal')
    
    doc.add_heading('Voorbeeld H3', level=3)
    doc.add_paragraph('Dit is normale body tekst (Kentledge Regular, 10.5 pt, kleur Ink). Gebruik dit voor alle lopende teksten, paragrafen en uitleg. De tekst is direct, menselijk, concreet en vooruitkijkend.', style='Normal')
    
    # Save the document
    output_path = r'C:\NOPPA\OneDrive - Noppa Solutions & Consultants\FRESH\Branding\Noppa-Basis-Sjabloon.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    create_template()
