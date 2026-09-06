import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def hex_to_rgb(hex_code):
    hex_code = hex_code.lstrip('#')
    return tuple(int(hex_code[i:i+2], 16) for i in (0, 2, 4))

def create_advanced():
    base_path = r'C:\NOPPA\OneDrive - Noppa Solutions & Consultants\FRESH\Branding\Noppa-Basis-Sjabloon.docx'
    doc = Document(base_path)

    # 1. Add Cover Page (Voorblad)
    # We will insert a paragraph at the very beginning of the document.
    cover_para = doc.paragraphs[0].insert_paragraph_before()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run()
    
    img_path = r'C:\NOPPA\OneDrive - Noppa Solutions & Consultants\FRESH\Branding\noppa-social\07-announcement-1200x1200.png'
    if os.path.exists(img_path):
        run.add_picture(img_path, width=Inches(6.0))
    
    title_para = doc.paragraphs[0].insert_paragraph_before('Document Titel', style='Heading 1')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para = doc.paragraphs[0].insert_paragraph_before('Subtitel of Datum', style='Heading 2')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a page break after the cover image
    page_break_para = doc.paragraphs[0].insert_paragraph_before()
    page_break_para.add_run().add_break()
    
    # Actually, inserting before paragraph[0] puts them in reverse order if we do it one by one.
    # Let's just create a new document, add the cover, then copy the contents, OR clear the old doc and recreate.
    # A better approach: We'll just build a brand new document from scratch, since the base one just has dummy text.
    pass

def build_new_doc():
    doc = Document()
    styles = doc.styles

    # Colors
    INK = '#1A2440'
    NAVY = '#0F2A66'
    ROYAL = '#2060E0'
    
    def setup_style(style, font_name, size_pt, color_hex, bold=False):
        font = style.font
        font.name = font_name
        font.size = Pt(size_pt)
        r, g, b = hex_to_rgb(color_hex)
        font.color.rgb = RGBColor(r, g, b)
        font.bold = bold

    # Setup styles
    setup_style(styles['Normal'], 'Kentledge', 10.5, INK, bold=False)
    setup_style(styles['Heading 1'], 'Kentledge', 36, NAVY, bold=True)
    setup_style(styles['Heading 2'], 'Kentledge', 22, NAVY, bold=True)
    setup_style(styles['Heading 3'], 'Kentledge', 14, NAVY, bold=True)
    
    # 1. Cover Page
    doc.add_heading('Document Titel', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Subtitel of Datum', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    img_path = r'C:\NOPPA\OneDrive - Noppa Solutions & Consultants\FRESH\Branding\noppa-social\07-announcement-1200x1200.png'
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Inches(5.0))
        
    doc.add_page_break()
    
    # 2. Header and Footer
    section = doc.sections[0]
    
    header = section.header
    h_p = header.paragraphs[0]
    h_p.text = "Noppa Solutions & Consultants"
    h_p.style = doc.styles['Normal']
    h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.text = "Boosting Business Productivity  |  www.noppa.nl  |  Pijlkruid 44, Berlicum"
    f_p.style = doc.styles['Normal']
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_p.runs[0].font.size = Pt(8.5)
    f_p.runs[0].font.color.rgb = RGBColor(*hex_to_rgb('#6B7280')) # Slate

    # 3. Content
    doc.add_heading('Inhoud', level=1)
    doc.add_paragraph('Dit sjabloon bevat een voorblad, headers, footers en tabelstijlen in de Noppa huisstijl.')
    
    doc.add_heading('Voorbeeld Tabel', level=2)
    
    # 4. Styled Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kolom 1'
    hdr_cells[1].text = 'Kolom 2'
    hdr_cells[2].text = 'Kolom 3'
    
    # Format header row (Background Royal, Text White)
    for cell in hdr_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), ROYAL))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                
    # Add some rows
    for i in range(3):
        row_cells = table.add_row().cells
        row_cells[0].text = f'Rij {i+1} Kolom 1'
        row_cells[1].text = f'Rij {i+1} Kolom 2'
        row_cells[2].text = f'Rij {i+1} Kolom 3'
        
        # Alternate row shading (Light: #F1F5FB)
        if i % 2 == 0:
            for cell in row_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="F1F5FB"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
    doc.add_paragraph('')
    doc.add_paragraph('Einde van het voorbeeld.', style='Normal')
    
    output_path = r'C:\NOPPA\OneDrive - Noppa Solutions & Consultants\FRESH\Branding\Noppa-Uitgebreid-Sjabloon.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    build_new_doc()
